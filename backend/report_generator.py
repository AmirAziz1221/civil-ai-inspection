import os
import io
from datetime import datetime
from pathlib import Path
from typing import Dict, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, inch
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        Image as RLImage, HRFlowable, PageBreak
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# ── Constants (Spanish) ────────────────────────────────────────────────────────

CLIENT_NAME  = "Dani"
REPORT_TITLE = "Informe de Detección de Patologías en Infraestructura Civil"
DISCLAIMER   = (
    "Este informe generado por IA es únicamente para apoyo preliminar de inspección. "
    "Todos los hallazgos y recomendaciones deben ser verificados y validados por un "
    "ingeniero civil cualificado y colegiado antes de tomar cualquier decisión de "
    "reparación o actuación. El sistema de detección por IA puede no identificar todos "
    "los defectos y no debe utilizarse como única base para decisiones de ingeniería "
    "o seguridad críticas."
)

SEVERITY_LABELS = {
    "Critical": "Crítico",
    "High":     "Alto",
    "Medium":   "Medio",
    "Low":      "Bajo",
}

SEVERITY_COLORS_HEX = {
    "Critical": "#DC2626",
    "High":     "#EA580C",
    "Medium":   "#D97706",
    "Low":      "#16A34A",
}


# ══════════════════════════════════════════════════════════════════════════════
#  BATCH — WORD REPORT
# ══════════════════════════════════════════════════════════════════════════════

def generate_batch_word_report(inspection: Dict, ai_report: Dict, output_path: str):
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(2.5)

    total_images = inspection.get("total_images", len(inspection.get("detections", [])))

    # Portada
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(REPORT_TITLE)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Preparado para: {CLIENT_NAME}").font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Inspeccion por Lotes - {total_images} Imagenes")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")

    doc.add_page_break()

    # 1. Informacion del Proyecto
    _add_heading(doc, "1. Informacion del Proyecto")
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("Fecha del Informe",   datetime.now().strftime("%d/%m/%Y %H:%M UTC")),
        ("Cliente",             CLIENT_NAME),
        ("Tipo de Activo",      inspection.get("asset_type", "N/A")),
        ("Modelo de Deteccion", inspection.get("model_name", "N/A")),
        ("Total de Imagenes",   str(total_images)),
    ]
    for i, (label, value) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()

    # 2. Resumen General
    _add_heading(doc, "2. Resumen General de Deteccion")
    total = sum(img.get("total_defects", 0) for img in inspection.get("detections", []))
    doc.add_paragraph(f"Total de imagenes analizadas: {len(inspection.get('detections', []))}")
    doc.add_paragraph(f"Total de defectos detectados: {total}")

    sev = inspection.get("severity_summary", {})
    sev_table = doc.add_table(rows=2, cols=4)
    sev_table.style = "Table Grid"
    for i, (eng, esp) in enumerate(SEVERITY_LABELS.items()):
        sev_table.rows[0].cells[i].text = esp
        sev_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        sev_table.rows[1].cells[i].text = str(sev.get(eng, 0))
    doc.add_paragraph()

    # 3. Resultados por Imagen
    _add_heading(doc, "3. Resultados por Imagen")
    for idx, img_result in enumerate(inspection.get("detections", []), 1):
        _add_heading(doc, f"3.{idx}  Imagen: {img_result.get('filename', f'Imagen {idx}')}", level=2)

        if img_result.get("error"):
            doc.add_paragraph(f"Error al procesar la imagen: {img_result['error']}")
            continue

        doc.add_paragraph(f"Total de defectos: {img_result.get('total_defects', 0)}")
        sev_val = SEVERITY_LABELS.get(img_result.get('overall_severity',''), img_result.get('overall_severity','N/A'))
        doc.add_paragraph(f"Gravedad general: {sev_val}")

        ann = img_result.get("annotated_image", "")
        if ann:
            local = Path(ann.lstrip("/"))
            if local.exists():
                try:
                    doc.add_picture(str(local), width=Inches(5))
                except Exception:
                    pass

        dets = img_result.get("detections", [])
        if dets:
            det_table = doc.add_table(rows=1, cols=4)
            det_table.style = "Table Grid"
            for i, h in enumerate(["#", "Clase de Defecto", "Confianza", "Gravedad"]):
                det_table.rows[0].cells[i].text = h
                det_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
            for i, d in enumerate(dets, 1):
                row = det_table.add_row().cells
                row[0].text = str(i)
                row[1].text = d.get("class", "").replace("_", " ").title()
                row[2].text = f"{d.get('confidence', 0):.1%}"
                row[3].text = SEVERITY_LABELS.get(d.get("severity",""), d.get("severity","N/A"))
        else:
            doc.add_paragraph("No se detectaron defectos en esta imagen.")
        doc.add_paragraph()

    _add_ai_sections(doc, ai_report, start_number=4)
    _add_disclaimer(doc)
    doc.save(output_path)


# ══════════════════════════════════════════════════════════════════════════════
#  SINGLE IMAGE — WORD REPORT
# ══════════════════════════════════════════════════════════════════════════════

def generate_word_report(inspection: Dict, ai_report: Dict, output_path: str):
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(2.5)

    _add_title_page(doc, inspection)
    doc.add_page_break()
    _add_project_info(doc, inspection)
    _add_detection_summary(doc, inspection)
    _add_visual_evidence(doc, inspection)
    _add_ai_sections(doc, ai_report)
    _add_disclaimer(doc)
    doc.save(output_path)


def _add_title_page(doc: Document, inspection: dict):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(REPORT_TITLE)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    doc.add_paragraph()

    items = [
        (f"Preparado para: {CLIENT_NAME}", 14, True),
        (f"Fecha: {datetime.now().strftime('%d/%m/%Y')}", 12, False),
        (f"Tipo de Activo: {inspection.get('asset_type', 'N/A')}", 12, False),
        (f"ID de Inspeccion: {inspection.get('id','N/A')[:8].upper()}", 10, False),
    ]
    for text, size, bold in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        if size == 10:
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def _add_project_info(doc: Document, inspection: dict):
    _add_heading(doc, "1. Informacion del Proyecto")
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("Fecha del Informe",   datetime.now().strftime("%d/%m/%Y %H:%M UTC")),
        ("Cliente",             CLIENT_NAME),
        ("Tipo de Activo",      inspection.get("asset_type", "N/A")),
        ("Modelo de Deteccion", inspection.get("model_name", "N/A")),
        ("Archivo de Imagen",   inspection.get("original_filename", "N/A")),
        ("ID de Inspeccion",    inspection.get("id", "N/A")[:8].upper()),
    ]
    for i, (label, value) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()


def _add_detection_summary(doc: Document, inspection: dict):
    _add_heading(doc, "2. Resumen de Deteccion")
    detections = inspection.get("detections", [])
    severity   = inspection.get("severity_summary", {})

    p = doc.add_paragraph()
    p.add_run(f"Total de defectos detectados: {len(detections)}").font.bold = True

    sev_table = doc.add_table(rows=2, cols=4)
    sev_table.style = "Table Grid"
    for i, (eng, esp) in enumerate(SEVERITY_LABELS.items()):
        sev_table.rows[0].cells[i].text = esp
        sev_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        sev_table.rows[1].cells[i].text = str(severity.get(eng, 0))
    doc.add_paragraph()

    if detections:
        _add_heading(doc, "Detalles de Deteccion", level=2)
        det_table = doc.add_table(rows=1, cols=4)
        det_table.style = "Table Grid"
        for i, h in enumerate(["#", "Clase de Defecto", "Confianza", "Gravedad"]):
            det_table.rows[0].cells[i].text = h
            det_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        for idx, det in enumerate(detections, 1):
            row = det_table.add_row().cells
            row[0].text = str(idx)
            row[1].text = det.get("class", "").replace("_", " ").title()
            row[2].text = f"{det.get('confidence', 0):.1%}"
            row[3].text = SEVERITY_LABELS.get(det.get("severity",""), det.get("severity","N/A"))
    doc.add_paragraph()


def _add_visual_evidence(doc: Document, inspection: dict):
    _add_heading(doc, "3. Evidencia Visual")
    upload_dir = Path("uploads")
    image_id   = inspection.get("image_id", "")
    orig_path  = None
    for f in upload_dir.iterdir():
        if f.stem == image_id:
            orig_path = f
            break

    ann_url  = inspection.get("annotated_image", "")
    ann_path = None
    if ann_url:
        local = Path(ann_url.lstrip("/"))
        if local.exists():
            ann_path = local

    for label, path in [
        ("Imagen Original", orig_path),
        ("Imagen Anotada (Deteccion IA)", ann_path),
    ]:
        if path and Path(path).exists():
            p = doc.add_paragraph(label)
            p.runs[0].font.bold = True
            try:
                doc.add_picture(str(path), width=Inches(5.5))
                doc.add_paragraph()
            except Exception:
                doc.add_paragraph(f"[Imagen: {path}]")


def _add_ai_sections(doc: Document, ai_report: dict, start_number: int = 4):
    sections = [
        ("Interpretacion de Ingenieria Civil por IA", "executive_summary"),
        ("Descripcion de Defectos",                   "defect_descriptions"),
        ("Posibles Causas",                           "possible_causes"),
        ("Evaluacion de Gravedad",                    "severity_assessment"),
        ("Evaluacion de Riesgos",                     "risk_explanation"),
        ("Acciones Recomendadas",                     "recommended_actions"),
        ("Nivel de Prioridad",                        "priority_level"),
        ("Conclusion Final",                          "final_conclusion"),
    ]
    for i, (heading, key) in enumerate(sections):
        _add_heading(doc, f"{start_number + i}. {heading}")
        content = ai_report.get(key, "No disponible.")
        doc.add_paragraph(content)
        doc.add_paragraph()


def _add_disclaimer(doc: Document):
    doc.add_page_break()
    _add_heading(doc, "Descargo de Responsabilidad")
    p = doc.add_paragraph(DISCLAIMER)
    p.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    p.runs[0].font.italic = True


def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)


# ══════════════════════════════════════════════════════════════════════════════
#  PDF REPORT
# ══════════════════════════════════════════════════════════════════════════════

def generate_pdf_report(inspection: Dict, ai_report: Dict, output_path: str):
    if not REPORTLAB_AVAILABLE:
        _generate_pdf_fallback(inspection, ai_report, output_path)
        return

    is_batch = inspection.get("is_batch", False)

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm
    )

    BLUE = colors.HexColor("#1E40AF")
    GRAY = colors.HexColor("#6B7280")
    BG   = colors.HexColor("#F8FAFC")

    styles = getSampleStyleSheet()
    title_style      = ParagraphStyle("T",  fontSize=20, textColor=BLUE, fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=12)
    h1_style         = ParagraphStyle("H1", fontSize=14, textColor=BLUE, fontName="Helvetica-Bold", spaceAfter=6, spaceBefore=12)
    h2_style         = ParagraphStyle("H2", fontSize=11, textColor=BLUE, fontName="Helvetica-Bold", spaceAfter=4, spaceBefore=8)
    body_style       = ParagraphStyle("Bo", fontSize=10, spaceAfter=6, leading=14, alignment=TA_JUSTIFY)
    center_style     = ParagraphStyle("Ce", fontSize=11, alignment=TA_CENTER, spaceAfter=4)
    disclaimer_style = ParagraphStyle("Di", fontSize=9, textColor=GRAY, fontName="Helvetica-Oblique", alignment=TA_JUSTIFY)

    def make_table(data, widths):
        t = Table(data, colWidths=widths)
        t.setStyle(TableStyle([
            ("BACKGROUND",     (0,0), (-1,0), BLUE),
            ("TEXTCOLOR",      (0,0), (-1,0), colors.white),
            ("FONTNAME",       (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTNAME",       (0,1), (0,-1), "Helvetica-Bold"),
            ("GRID",           (0,0), (-1,-1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, BG]),
            ("FONTSIZE",       (0,0), (-1,-1), 9),
            ("PADDING",        (0,0), (-1,-1), 6),
        ]))
        return t

    story = []

    # Portada
    story.append(Spacer(1, 2*cm))
    story.append(Paragraph(REPORT_TITLE, title_style))
    story.append(Spacer(1, 0.5*cm))
    story.append(HRFlowable(width="100%", thickness=2, color=BLUE))
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph(f"Preparado para: <b>{CLIENT_NAME}</b>", center_style))
    if is_batch:
        total_images = inspection.get("total_images", len(inspection.get("detections", [])))
        story.append(Paragraph(f"Inspeccion por Lotes - {total_images} Imagenes", center_style))
    story.append(Paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}", center_style))
    story.append(Paragraph(f"Tipo de Activo: {inspection.get('asset_type', 'N/A')}", center_style))
    story.append(Spacer(1, 2*cm))
    story.append(PageBreak())

    # 1. Informacion del Proyecto
    story.append(Paragraph("1. Informacion del Proyecto", h1_style))
    info_data = [["Campo", "Valor"],
        ["Fecha del Informe",   datetime.now().strftime("%d/%m/%Y %H:%M UTC")],
        ["Cliente",             CLIENT_NAME],
        ["Tipo de Activo",      inspection.get("asset_type", "N/A")],
        ["Modelo de Deteccion", inspection.get("model_name", "N/A")],
    ]
    if is_batch:
        info_data.append(["Total de Imagenes", str(inspection.get("total_images", "N/A"))])
    else:
        info_data.append(["Archivo de Imagen", inspection.get("original_filename", "N/A")])
    info_data.append(["ID de Inspeccion", inspection.get("id", "N/A")[:8].upper()])
    story.append(make_table(info_data, [5*cm, 12*cm]))
    story.append(Spacer(1, 0.5*cm))

    # 2. Resumen de Deteccion
    story.append(Paragraph("2. Resumen de Deteccion", h1_style))
    total = inspection.get("total_defects", 0)
    story.append(Paragraph(f"<b>Total de defectos detectados: {total}</b>", body_style))

    sev = inspection.get("severity_summary", {})
    sev_data = [
        [SEVERITY_LABELS["Critical"], SEVERITY_LABELS["High"],
         SEVERITY_LABELS["Medium"],   SEVERITY_LABELS["Low"]],
        [str(sev.get("Critical",0)),  str(sev.get("High",0)),
         str(sev.get("Medium",0)),    str(sev.get("Low",0))]
    ]
    sev_t = Table(sev_data, colWidths=[4.25*cm]*4)
    sev_t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (0,0), colors.HexColor("#DC2626")),
        ("BACKGROUND", (1,0), (1,0), colors.HexColor("#EA580C")),
        ("BACKGROUND", (2,0), (2,0), colors.HexColor("#D97706")),
        ("BACKGROUND", (3,0), (3,0), colors.HexColor("#16A34A")),
        ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
        ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
        ("ALIGN",      (0,0), (-1,-1), "CENTER"),
        ("GRID",       (0,0), (-1,-1), 0.5, colors.grey),
        ("FONTSIZE",   (0,0), (-1,-1), 10),
        ("PADDING",    (0,0), (-1,-1), 8),
    ]))
    story.append(sev_t)
    story.append(Spacer(1, 0.3*cm))

    # 3. Evidencia / Resultados
    if is_batch:
        story.append(Paragraph("3. Resultados por Imagen", h1_style))
        for idx, img in enumerate(inspection.get("detections", []), 1):
            story.append(Paragraph(f"3.{idx}  {img.get('filename', f'Imagen {idx}')}", h2_style))
            if img.get("error"):
                story.append(Paragraph(f"Error: {img['error']}", body_style))
                continue
            sev_label = SEVERITY_LABELS.get(img.get("overall_severity",""), img.get("overall_severity","N/A"))
            story.append(Paragraph(
                f"Defectos: <b>{img.get('total_defects',0)}</b>  |  Gravedad: <b>{sev_label}</b>",
                body_style))
            ann = img.get("annotated_image","")
            if ann:
                local = Path(ann.lstrip("/"))
                if local.exists():
                    try:
                        story.append(RLImage(str(local), width=14*cm, height=8*cm, kind="proportional"))
                        story.append(Spacer(1, 0.2*cm))
                    except Exception:
                        pass
            dets = img.get("detections", [])
            if dets:
                det_data = [["#", "Clase de Defecto", "Confianza", "Gravedad"]]
                for i, d in enumerate(dets, 1):
                    det_data.append([str(i),
                        d.get("class","").replace("_"," ").title(),
                        f"{d.get('confidence',0):.1%}",
                        SEVERITY_LABELS.get(d.get("severity",""), d.get("severity","N/A"))])
                story.append(make_table(det_data, [1*cm, 7*cm, 4*cm, 5*cm]))
            else:
                story.append(Paragraph("No se detectaron defectos en esta imagen.", body_style))
            story.append(Spacer(1, 0.4*cm))
    else:
        story.append(Paragraph("3. Evidencia Visual", h1_style))
        image_id   = inspection.get("image_id", "")
        upload_dir = Path("uploads")
        orig_path  = None
        for f in upload_dir.iterdir():
            if f.stem == image_id:
                orig_path = f
                break
        ann_url  = inspection.get("annotated_image", "")
        ann_path = Path(ann_url.lstrip("/")) if ann_url else None
        for label, path in [("Imagen Original", orig_path), ("Imagen Anotada (Deteccion IA)", ann_path)]:
            if path and Path(path).exists():
                story.append(Paragraph(f"<b>{label}</b>", body_style))
                try:
                    story.append(RLImage(str(path), width=15*cm, height=9*cm, kind="proportional"))
                    story.append(Spacer(1, 0.3*cm))
                except Exception:
                    story.append(Paragraph(f"[Imagen no disponible: {path}]", body_style))
        dets = inspection.get("detections", [])
        if dets:
            story.append(Paragraph("Detalles de Deteccion", h2_style))
            det_data = [["#", "Clase de Defecto", "Confianza", "Gravedad"]]
            for i, d in enumerate(dets, 1):
                det_data.append([str(i),
                    d.get("class","").replace("_"," ").title(),
                    f"{d.get('confidence',0):.1%}",
                    SEVERITY_LABELS.get(d.get("severity",""), d.get("severity","N/A"))])
            story.append(make_table(det_data, [1*cm, 7*cm, 4*cm, 5*cm]))
        story.append(Spacer(1, 0.5*cm))

    # Secciones de IA
    ai_sections = [
        ("4. Interpretacion de Ingenieria Civil por IA", "executive_summary"),
        ("5. Descripcion de Defectos",                   "defect_descriptions"),
        ("6. Posibles Causas",                           "possible_causes"),
        ("7. Evaluacion de Gravedad",                    "severity_assessment"),
        ("8. Evaluacion de Riesgos",                     "risk_explanation"),
        ("9. Acciones Recomendadas",                     "recommended_actions"),
        ("10. Nivel de Prioridad",                       "priority_level"),
        ("11. Conclusion Final",                         "final_conclusion"),
    ]
    for heading, key in ai_sections:
        story.append(Paragraph(heading, h1_style))
        content = ai_report.get(key, "No disponible.")
        story.append(Paragraph(content, body_style))
        story.append(Spacer(1, 0.3*cm))

    # Descargo
    story.append(PageBreak())
    story.append(Paragraph("Descargo de Responsabilidad", h1_style))
    story.append(Paragraph(DISCLAIMER, disclaimer_style))

    doc.build(story)


def _generate_pdf_fallback(inspection: dict, ai_report: dict, output_path: str):
    try:
        import subprocess
        html = _build_html_report(inspection, ai_report)
        html_path = output_path.replace(".pdf", "_temp.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html)
        subprocess.run(["wkhtmltopdf", html_path, output_path], check=True, capture_output=True)
        os.remove(html_path)
    except Exception:
        with open(output_path, "wb") as f:
            f.write(b"%PDF-1.4\n%Requiere ReportLab\n")


def _build_html_report(inspection: dict, ai_report: dict) -> str:
    return f"""<!DOCTYPE html><html><head><meta charset='utf-8'>
    <title>{REPORT_TITLE}</title></head><body>
    <h1>{REPORT_TITLE}</h1>
    <p>Preparado para: {CLIENT_NAME}</p>
    <p>Fecha: {datetime.now().strftime('%d/%m/%Y')}</p>
    <h2>Resumen</h2>
    <p>{ai_report.get('executive_summary', 'N/A')}</p>
    <p><em>{DISCLAIMER}</em></p>
    </body></html>"""